# -*- coding: utf-8 -*-
"""
generator.py - Word document generator for calculation reports.

Uses `pandas`, `python-docx`, and `matplotlib` to generate Word documents
that contain structural report tables and charts.

Usage:
    from report import ReportGenerator

    rg = ReportGenerator(model)
    rg.generate("calculation_report.docx", frame_names=["56", "57", "59"])
"""

import os
import tempfile
from typing import List, Optional
from lxml import etree

import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .data_collector import (
    collect_frame_element_info,
    collect_steel_design_results,
    FrameElementInfo,
    SteelDesignResultRow,
)

# Word XML namespace
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ==================== Word table helpers ====================

def _set_cell_text(cell, text: str, bold: bool = False, size: float = 9):
    """Set cell text and center it."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    spacing = parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'
    )
    pPr.append(spacing)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.bold = bold
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))


def _set_table_borders(table):
    """Apply borders to the whole table."""
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblPr.append(parse_xml(borders_xml))


def _shade_header_row(row):
    """Shade the header row gray."""
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)


def _add_title(doc: Document, text: str):
    """Add a centered bold title."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.bold = True


def _df_to_table(doc: Document, df: pd.DataFrame, headers: List[str]):
    """Write a DataFrame into a Word table."""
    num_rows = len(df) + 1
    num_cols = len(headers)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    header_row = table.rows[0]
    _shade_header_row(header_row)
    for i, h in enumerate(headers):
        _set_cell_text(header_row.cells[i], h, bold=True, size=8)
    for r_idx, (_, row_data) in enumerate(df.iterrows(), start=1):
        row = table.rows[r_idx]
        for c_idx in range(num_cols):
            _set_cell_text(row.cells[c_idx], str(row_data.iloc[c_idx]), size=9)
    return table


# ==================== DataFrame builders ====================

def _build_element_info_df(data: List[FrameElementInfo]) -> pd.DataFrame:
    """Convert element info rows to a DataFrame."""
    rows = []
    for info in data:
        rows.append({
            "No.": info.no,
            "Section": info.section_name,
            "Length": f"{info.length:.3f}",
            "Unbraced\nLength Ratio\n(L2)": f"{info.unbraced_ratio_major:.3f}",
            "Unbraced\nLength Ratio\n(L3)": f"{info.unbraced_ratio_minor:.3f}",
            "Effective\nLength Factor\n(mu2)": f"{info.mue_major:.3f}",
            "Effective\nLength Factor\n(mu3)": f"{info.mue_minor:.3f}",
            "I-End Release": info.release_i,
            "J-End Release": info.release_j,
        })
    return pd.DataFrame(rows)


def _fmt(val: float) -> str:
    """Format values, showing `-` when `0` means not applicable."""
    if val == 0.0:
        return "—"
    return f"{val:.3f}"


def _build_steel_design_df(data: List[SteelDesignResultRow]) -> pd.DataFrame:
    """Convert steel design result rows to a DataFrame."""
    rows = []
    for r in data:
        rows.append({
            "No.": r.index,
            "Frame": r.frame_name,
            "Stress Ratio": f"{r.ratio:.3f}",
            "N(kN)": f"{r.n:.3f}",
            "M3(kN·m)": f"{r.m3:.3f}",
            "M2(kN·m)": f"{r.m2:.3f}",
            "φ3": _fmt(r.phi3),
            "φ2": _fmt(r.phi2),
            "φb3": _fmt(r.phi_b3),
            "φb2": _fmt(r.phi_b2),
            "Slenderness\nMajor": _fmt(r.slenderness_major),
            "Slenderness\nMinor": _fmt(r.slenderness_minor),
            "Exceeded": r.exceeded,
        })
    return pd.DataFrame(rows)


# ==================== Chart generation ====================

def _generate_bar_chart(design_data: List[SteelDesignResultRow]) -> str:
    """
    Generate a bar chart of frame stress ratios (Figure 5.1).

    Returns:
        Temporary image file path.
    """
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial"]
    plt.rcParams["axes.unicode_minus"] = False

    names = [r.frame_name for r in design_data]
    ratios = [r.ratio for r in design_data]

    fig, ax = plt.subplots(figsize=(10, 4))
    ax.bar(range(len(names)), ratios, color="blue", edgecolor="blue", width=0.7)
    ax.set_xlabel("Frames", fontsize=12)
    ax.set_ylabel("Stress Ratio", fontsize=12)
    ax.set_xticks([])
    ax.grid(axis="both", linestyle="--", alpha=0.4, color="gray")
    ax.set_xlim(-0.5, len(names) - 0.5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()

    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=200)
    plt.close(fig)
    return tmp.name


def _generate_pie_chart(design_data: List[SteelDesignResultRow]) -> str:
    """
    Generate a pie chart of stress-ratio distribution (Figure 5.2).

    Ranges: `<=0.5`, `0.5~0.7`, `0.7~0.9`, `0.9~1.0`, `>1.0`
    """
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial"]
    plt.rcParams["axes.unicode_minus"] = False

    bin_defs = [
        ("<=0.5", 0.0, 0.5),
        ("0.5~0.7", 0.5, 0.7),
        ("0.7~0.9", 0.7, 0.9),
        ("0.9~1.0", 0.9, 1.0),
        (">1.0", 1.0, float("inf")),
    ]
    all_colors = ["#00FFFF", "#7FFF00", "#FFD700", "#FF8C00", "#FF0000"]

    total = len(design_data)
    counts = []
    labels = []
    colors = []
    for i, (label, lo, hi) in enumerate(bin_defs):
        if hi == float("inf"):
            c = sum(1 for r in design_data if r.ratio > lo)
        else:
            c = sum(1 for r in design_data if lo < r.ratio <= hi)
            if lo == 0.0:
                c = sum(1 for r in design_data if r.ratio <= hi)
        if c > 0:
            pct = c / total * 100
            labels.append(f"{label}\n({pct:.2f}%)")
            counts.append(c)
            colors.append(all_colors[i])

    fig, ax = plt.subplots(figsize=(6, 6))
    ax.pie(
        counts, labels=labels, colors=colors,
        startangle=90, counterclock=False,
        textprops={"fontsize": 10},
    )
    ax.axis("equal")
    fig.tight_layout()

    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=200)
    plt.close(fig)
    return tmp.name


def _add_chart_with_caption(doc: Document, image_path: str, caption: str, width: float = 5.5):
    """Insert an image and add a centered caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(12)
    r.font.name = "SimSun"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


# ==================== Main generator ====================

class ReportGenerator:
    """
    Calculation report generator.

    Collects data from a SAP2000 model and generates a Word report.

    Args:
        model: SAP2000 SapModel object

    Usage:
        rg = ReportGenerator(model)
        rg.generate("calculation_report.docx", frame_names=["56", "57", "59"])
    """

    def __init__(self, model):
        self._model = model

    def generate(
        self,
        output_path: str,
        frame_names: List[str] = None,
        element_table_number: str = "2.5",
        design_table_number: str = "5.1",
        include_element_info: bool = True,
        include_design_results: bool = True,
        include_charts: bool = True,
    ) -> str:
        """
        Generate a Word calculation report.

        Args:
            output_path: Output `.docx` path
            frame_names: Frame names; uses all model frames when `None`
            element_table_number: Element information table number
            design_table_number: Design result table number
            include_element_info: Whether to include the element info table
            include_design_results: Whether to include the design result table
            include_charts: Whether to include stress ratio charts

        Returns:
            Absolute output file path.
        """
        from structure_core.frame import Frame

        if frame_names is None:
            frame_names = Frame.get_name_list(self._model)

        doc = Document()

        # Set default font.
        style = doc.styles["Normal"]
        style.font.name = "SimSun"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

        design_data = None

        if include_element_info:
            element_data = collect_frame_element_info(self._model, frame_names)
            df = _build_element_info_df(element_data)
            _add_title(doc, f"Table {element_table_number} Element Information")
            _df_to_table(doc, df, list(df.columns))
            doc.add_paragraph()

        if include_design_results or include_charts:
            design_data = collect_steel_design_results(self._model, frame_names)

        if include_design_results and design_data:
            df = _build_steel_design_df(design_data)
            _add_title(doc, f"Table {design_table_number} Steel Design Results")
            _df_to_table(doc, df, list(df.columns))
            doc.add_paragraph()

        # Charts
        if include_charts and design_data:
            tmp_files = []
            try:
                # Figure 5.1 bar chart
                bar_path = _generate_bar_chart(design_data)
                tmp_files.append(bar_path)
                _add_chart_with_caption(
                    doc, bar_path,
                    f"Figure {design_table_number} Stress Ratio Distribution"
                )
                doc.add_paragraph()

                # Figure 5.2 pie chart
                pie_path = _generate_pie_chart(design_data)
                tmp_files.append(pie_path)
                _add_chart_with_caption(
                    doc, pie_path,
                    f"Figure {design_table_number.split('.')[0]}.2 Stress Ratio Share Distribution",
                    width=4.0,
                )
            finally:
                # Clean up temporary files.
                for f in tmp_files:
                    try:
                        os.unlink(f)
                    except OSError:
                        pass

        doc.save(output_path)
        return os.path.abspath(output_path)

    def generate_element_info_only(
        self,
        output_path: str,
        frame_names: List[str] = None,
        table_number: str = "2.5",
    ) -> str:
        """Generate only the element information table."""
        return self.generate(
            output_path,
            frame_names=frame_names,
            element_table_number=table_number,
            include_element_info=True,
            include_design_results=False,
            include_charts=False,
        )

    def generate_design_results_only(
        self,
        output_path: str,
        frame_names: List[str] = None,
        table_number: str = "5.1",
    ) -> str:
        """Generate only the steel design result table."""
        return self.generate(
            output_path,
            frame_names=frame_names,
            design_table_number=table_number,
            include_element_info=False,
            include_design_results=True,
            include_charts=True,
        )
    